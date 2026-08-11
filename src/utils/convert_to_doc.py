#!/usr/bin/env python3
"""
将Markdown文档转换为DOC格式的脚本
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import re

def setup_chinese_font(doc):
    """设置中文字体支持"""
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    # 设置中文字体
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_paragraph(doc, text, style=None):
    """添加段落"""
    para = doc.add_paragraph(text, style=style)
    # 设置中文字体
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para

def add_code_block(doc, code):
    """添加代码块"""
    para = doc.add_paragraph(code)
    para.style = 'NoSpacing'
    # 设置等宽字体和背景色
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    # 设置段落格式
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    return para

def parse_markdown_to_doc(markdown_file, output_file):
    """将Markdown文件转换为DOC文件"""
    doc = Document()

    # 设置中文字体
    setup_chinese_font(doc)

    # 读取Markdown文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 处理标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            title_text = line.lstrip('#').strip()
            add_heading(doc, title_text, min(level, 3))
            i += 1
            continue

        # 处理列表
        if line.startswith('- ') or line.startswith('* '):
            text = line.lstrip('-*').strip()
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # 处理编号列表
        if re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.', '', line).strip()
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 处理表格（简化处理）
        if line.startswith('|') and '|' in line:
            # 跳过表格分隔线
            if set(line.strip()) == {'|', '-', ' '}:
                i += 1
                continue

            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_row = doc.add_table(1, len(cells))
            table_row.autofit = True
            for j, cell in enumerate(cells):
                table_row.rows[0].cells[j].text = cell
                # 设置单元格字体
                for paragraph in table_row.rows[0].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # 处理代码行
        if line.startswith('```') or (line.startswith('    ') and not line.strip() == ''):
            code_text = line.replace('```', '').lstrip()
            add_code_block(doc, code_text)
            i += 1
            continue

        # 处理普通段落
        if line.strip():
            # 处理粗体
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # 处理斜体
            line = re.sub(r'\*(.*?)\*', r'\1', line)
            add_paragraph(doc, line)

        i += 1

    # 保存文档
    doc.save(output_file)
    print(f"文档已成功转换为DOC格式: {output_file}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        markdown_file = sys.argv[1]
    else:
        markdown_file = "TensorRT_环境配置报告.md"

    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        output_file = "TensorRT_环境配置报告.docx"

    try:
        parse_markdown_to_doc(markdown_file, output_file)
    except FileNotFoundError:
        print(f"错误: 找不到文件 {markdown_file}")
        print("请确保文件路径正确")
        sys.exit(1)
    except Exception as e:
        print(f"转换失败: {e}")
        sys.exit(1)
